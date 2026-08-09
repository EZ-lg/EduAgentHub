"""
文档解析 — 从上传文件中提取纯文本（知识库入库的前置步骤）

支持格式：
- PDF      （PyPDF2）
- Word     （python-docx，仅 .docx；旧版 .doc 不支持，明确报错）
- Markdown / TXT （utf-8，GBK 兜底）

对外只暴露 parse_file(file_path, file_type)，返回 (text, error)：
- 成功：text 为纯文本，error 为 None
- 失败：text 为空串，error 为可读错误信息（路由层据此抛 400）
"""
from fastapi import HTTPException

# 支持的扩展名 → 格式标识
SUPPORTED_EXTS = {"pdf", "docx", "txt", "md", "markdown"}
# 扩展名归一化（.md / .markdown 统一为 md）
EXT_ALIAS = {"markdown": "md"}

# 单文件大小上限（50MB，需求 F9.1）
MAX_FILE_SIZE = 50 * 1024 * 1024


def parse_file(file_path: str, file_type: str):
    """解析文档 → (text, error)。file_type 为归一化格式（pdf/docx/txt/md）。"""
    if file_type not in SUPPORTED_EXTS and file_type not in EXT_ALIAS:
        return "", f"不支持的文件格式：{file_type or '未知'}（支持 PDF / Word / TXT / Markdown）"
    fmt = EXT_ALIAS.get(file_type, file_type)

    try:
        with open(file_path, "rb") as f:
            size = len(f.read())
        if size > MAX_FILE_SIZE:
            return "", f"文件超过 50MB 上限（当前 {size // 1024 // 1024}MB）"
    except OSError as e:
        return "", f"无法读取文件：{e}"

    try:
        if fmt == "pdf":
            return _parse_pdf(file_path), None
        if fmt == "docx":
            return _parse_docx(file_path), None
        return _parse_text(file_path), None  # txt / md
    except HTTPException:
        raise
    except Exception as e:
        return "", f"解析失败：{e}"


def _parse_pdf(file_path: str) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append(text)
    return "\n\n".join(pages).strip()


def _parse_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    # 表格内容也纳入（常见于课程大纲、收费表）
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _parse_text(file_path: str) -> str:
    """txt/md：utf-8 优先，GBK 兜底（大陆常见编码）"""
    for encoding in ("utf-8", "gbk", "utf-8-sig"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
    # 全部失败：读二进制按 utf-8 忽略错误，避免直接报错
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


def parse_text_string(content: bytes, file_type: str):
    """解析内存内容（供测试/预览复用），签名同 parse_file。"""
    fmt = EXT_ALIAS.get(file_type, file_type)
    if fmt == "pdf":
        from PyPDF2 import PdfReader
        from io import BytesIO

        reader = PdfReader(BytesIO(content))
        return "\n\n".join((p.extract_text() or "") for p in reader.pages).strip()
    if fmt == "docx":
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(content))
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()).strip()
    for encoding in ("utf-8", "gbk", "utf-8-sig"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore").strip()
