"""
P5 docx 导出 — 按机构模板《陈思敏 冲刺600分学习计划》排版生成单科学习计划 docx

模板排版规范（从模板 docx 提取）：
- 主标题：微软雅黑 22pt #1F4E79 加粗（居中）
- 副标题：微软雅黑 12pt #555555（居中）
- 方法总纲（红色callout）：宋体 10pt #C0392B 加粗
- 章节标题：微软雅黑 15pt #1F4E79 加粗；正文：宋体 11pt、1.5 倍行距
- 课程规划表格：Table Grid，表头微软雅黑 10pt 白字 #1F4E79 蓝底
- 写在最后：宋体 10.5pt #C0392B 加粗
"""
import json
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xC0, 0x39, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_font(run, name, size, color, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 中文需同时设置 eastAsia 字体，否则 Word 可能回退默认中文字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)


def _add_para(doc, text, name='宋体', size=11, color=None, bold=False,
              align=None, space_after=6, line_spacing=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    run = p.add_run(text)
    _set_font(run, name, size, color, bold)
    return p


def _set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _add_plan_table(doc, plan_rows):
    """课程规划表格：课时/课程内容/课时数/教师/时间安排/备注"""
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['课时', '课程内容', '课时数', '教师', '时间安排', '备注']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        _set_font(run, '微软雅黑', 10, WHITE, True)
        _set_cell_bg(cell, '1F4E79')
    for row in plan_rows:
        vals = [
            str(row.get('lesson') or ''),
            str(row.get('content') or ''),
            str(row.get('hours') if row.get('hours') is not None else ''),
            str(row.get('teacher_name') or ''),
            str(row.get('schedule') or ''),
            str(row.get('notes') or ''),
        ]
        cells = table.add_row().cells
        for i, v in enumerate(vals):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(v)
            _set_font(run, '宋体', 10, None, False)
    return table


def build_report_docx(report, subject=None, student=None, org_name=''):
    """根据 Report 生成单科学习计划 docx，返回 BytesIO"""
    try:
        content = json.loads(report.content_json or '{}')
    except (json.JSONDecodeError, TypeError):
        content = {}
    if not isinstance(content, dict):
        content = {}
    chapters = content.get('chapters') if isinstance(content.get('chapters'), list) else []
    plan = content.get('plan') if isinstance(content.get('plan'), list) else []

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(11)

    # 主标题（优先 content.title，带「·」样式）
    title = content.get('title') or report.title or '学习计划'
    _add_para(doc, title, name='微软雅黑', size=22, color=DARK_BLUE, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.2)

    # 副标题
    subtitle = content.get('subtitle') or ''
    if subtitle:
        _add_para(doc, subtitle, name='微软雅黑', size=12, color=GRAY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line_spacing=1.2)

    # 方法总纲（红色 callout）
    if content.get('summary'):
        _add_para(doc, content['summary'], name='宋体', size=10, color=RED, bold=True, space_after=10)

    # 章节
    for ch in chapters:
        if not isinstance(ch, dict):
            continue
        c_title = str(ch.get('title') or '').strip()
        c_body = str(ch.get('content') or '').strip()
        if not c_title and not c_body:
            continue
        if c_title:
            _add_para(doc, c_title, name='微软雅黑', size=15, color=DARK_BLUE, bold=True,
                      space_after=4, line_spacing=1.2)
        if c_body:
            _add_para(doc, c_body, name='宋体', size=11, space_after=8)
        # 「六、本学科落地规划」后插入课程规划表格
        if c_title[:2] == '六、' and plan:
            _add_plan_table(doc, plan)
            doc.add_paragraph()

    # 写在最后
    if content.get('conclusion'):
        _add_para(doc, content['conclusion'], name='宋体', size=10.5, color=RED, bold=True, space_after=6)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
